import logging
from docx import Document
from app import db, bedrock
from models import TranslationJob
import json
import os

logger = logging.getLogger(__name__)

def process_document(job_id):
    """Process and translate a document"""
    job = TranslationJob.query.get(job_id)
    if not job:
        logger.error(f"Job {job_id} not found")
        return

    try:
        job.status = 'processing'
        db.session.commit()

        # Load document
        doc = Document(job.file_path)
        total_paragraphs = len(doc.paragraphs)
        translated_paragraphs = []

        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                translated_paragraphs.append(paragraph.text)
                continue

            # Translate text using Bedrock
            translated_text = translate_text(
                paragraph.text,
                job.source_language,
                job.target_language
            )
            translated_paragraphs.append(translated_text)

            # Update progress
            progress = (i + 1) / total_paragraphs * 100
            job.progress = progress
            db.session.commit()

        # Create new document with translations
        new_doc = Document()
        for text in translated_paragraphs:
            new_doc.add_paragraph(text)

        # Save translated document
        translated_path = f"temp/translated_{job.original_filename}"
        new_doc.save(translated_path)
        job.translated_file_path = translated_path
        job.status = 'completed'
        job.message = 'Translation completed successfully'

    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        job.status = 'failed'
        job.message = f'Translation failed: {str(e)}'

    finally:
        db.session.commit()

def translate_text(text, source_lang, target_lang):
    """Translate text using AWS Bedrock"""
    prompt = f"""Translate the following text from {source_lang} to {target_lang}:
    "{text}"
    Translation:"""

    try:
        response = bedrock.invoke_model(
            modelId="anthropic.claude-3-sonnet-20240229-v1:0",
            body=json.dumps({
                "prompt": prompt,
                "max_tokens": 4096,
                "temperature": 0.1
            })
        )
        
        response_body = json.loads(response['body'].read())
        return response_body['completion'].strip()

    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        raise
